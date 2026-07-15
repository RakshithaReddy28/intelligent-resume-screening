from docx import Document
from app.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):
    """Parser for DOCX resume files."""

    def parse(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        self.validate_file(file_path)

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)
